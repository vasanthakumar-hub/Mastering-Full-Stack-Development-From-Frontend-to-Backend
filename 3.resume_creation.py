from docx import Document
from docx.shared import Pt

def create_resume():
    # Create a new Document
    doc = Document()

    # Add Title
    doc.add_heading('John Doe', 0)

    # Add Contact Information
    doc.add_paragraph('Email: johndoe@example.com\nPhone: +123456789\nLinkedIn: linkedin.com/in/johndoe\nGitHub: github.com/johndoe')

    # Add Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph(
        'Bachelor of Technology in Computer Science\n'
        'XYZ University, City, Country\n'
        'Graduation Year: 2024\n'
        'GPA: 3.8/4.0'
    )

    # Add Skills
    doc.add_heading('Skills', level=1)
    skills = doc.add_paragraph()
    skills.add_run('Programming Languages: ').bold = True
    skills.add_run('Python, JavaScript, HTML, CSS, SQL\n')
    skills.add_run('Web Technologies: ').bold = True
    skills.add_run('React, Node.js, Express.js, Django\n')
    skills.add_run('Tools: ').bold = True
    skills.add_run('Git, Docker, Jenkins\n')

    # Add Projects
    doc.add_heading('Projects', level=1)
    project1 = doc.add_paragraph()
    project1.add_run('Portfolio Website - ').bold = True
    project1.add_run('Developed a personal portfolio website using HTML, CSS, and JavaScript to showcase my projects and skills.')
    
    project2 = doc.add_paragraph()
    project2.add_run('E-commerce Website - ').bold = True
    project2.add_run('Built a full-stack e-commerce application using React and Node.js with features such as product listing, shopping cart, and payment integration.')

    # Add Experience (if any)
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Intern, ABC Company\n'
                      'June 2023 - August 2023\n'
                      '• Developed web applications using React and Node.js\n'
                      '• Collaborated with designers and back-end developers to create a seamless user experience\n'
                      '• Participated in code reviews and contributed to the improvement of code quality')

    # Add Certifications (if any)
    doc.add_heading('Certifications', level=1)
    doc.add_paragraph('Certified Web Developer, FreeCodeCamp\n'
                      'JavaScript Algorithms and Data Structures Certification, FreeCodeCamp')

    # Add Hobbies (optional)
    doc.add_heading('Hobbies', level=1)
    doc.add_paragraph('• Coding\n'
                      '• Blogging about web development\n'
                      '• Playing chess\n'
                      '• Traveling')

    # Save the document
    doc.save('John_Doe_Resume.docx')

if __name__ == "__main__":
    create_resume()